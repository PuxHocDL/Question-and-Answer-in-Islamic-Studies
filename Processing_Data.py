import os
import subprocess
from docx import Document
import re
import shutil

def convert_doc_to_docx(file_path, temp_dir="temp"):
    """Chuyển đổi file .doc sang .docx bằng LibreOffice"""
    try:
        if not os.path.exists(temp_dir):
            os.makedirs(temp_dir)
        output_path = os.path.join(temp_dir, f"{os.path.splitext(os.path.basename(file_path))[0]}.docx")
        subprocess.run([
            "libreoffice", "--headless", "--convert-to", "docx",
            file_path, "--outdir", temp_dir
        ], check=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE, text=True)
        if os.path.exists(output_path):
            print(f"Đã chuyển đổi {file_path} sang {output_path}")
            return output_path
        else:
            raise Exception("Chuyển đổi thất bại, file .docx không được tạo")
    except subprocess.CalledProcessError as e:
        print(f"Lỗi khi chuyển đổi {file_path}: {e.stderr}")
        return None
    except Exception as e:
        print(f"Lỗi khi chuyển đổi {file_path}: {e}")
        return None

def process_arabic_word_file(file_path, full_text):
    try:
        # Mở file Word
        doc = Document(file_path)
        
        # Thêm metadata file
        full_text.append(f"[FILE: {os.path.basename(file_path)}]\n")
        current_section = None
        
        # Đọc từng đoạn văn
        paragraph_buffer = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                # Kiểm tra tiêu đề
                if para.style.name.startswith('Heading') or re.match(r'^(النَّوْعُ|المجلد|مقدمة)', text):
                    # Lưu đoạn văn trước đó (nếu có)
                    if paragraph_buffer:
                        full_text.append("".join(paragraph_buffer) + "\n")
                        paragraph_buffer = []
                    current_section = text
                    full_text.append(f"[SECTION: {current_section}]\n")
                else:
                    # Thêm vào buffer để giữ đoạn văn liên tục
                    paragraph_buffer.append(text + " ")
        
        # Lưu đoạn văn cuối cùng
        if paragraph_buffer:
            full_text.append("".join(paragraph_buffer) + "\n")
        
        # Đọc bảng (nếu có)
        for table in doc.tables:
            full_text.append("[TABLE]\n")
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells]
                full_text.append("\t".join(row_text) + "\n")
        
        return True
    
    except Exception as e:
        print(f"Lỗi khi xử lý file {file_path}: {e}")
        full_text.append(f"[ERROR: {file_path}]\nLỗi: {e}\n")
        return False

def process_all_word_files(directory, output_file="all_output.txt"):
    # Kiểm tra xem thư mục có tồn tại không
    if not os.path.isdir(directory):
        print(f"Thư mục {directory} không tồn tại!")
        return
    
    # Danh sách lưu trữ toàn bộ nội dung
    full_text = []
    
    # Duyệt qua tất cả các file trong thư mục
    for filename in os.listdir(directory):
        if filename.lower().endswith(('.docx', '.doc')):
            file_path = os.path.join(directory, filename)
            print(f"Đang xử lý file: {file_path}")
            
            # Thử mở file với python-docx (chỉ hỗ trợ .docx)
            if filename.lower().endswith('.docx'):
                success = process_arabic_word_file(file_path, full_text)
                if not success:
                    print(f"Thử chuyển đổi {file_path} như file .doc...")
                    converted_path = convert_doc_to_docx(file_path)
                    if converted_path:
                        process_arabic_word_file(converted_path, full_text)
            else:
                print(f"File {filename} là định dạng .doc, đang chuyển đổi sang .docx...")
                converted_path = convert_doc_to_docx(file_path)
                if converted_path:
                    process_arabic_word_file(converted_path, full_text)
                else:
                    full_text.append(f"[ERROR: {filename}]\nLỗi: Không thể chuyển đổi file .doc sang .docx\n")
    
    # Lưu toàn bộ nội dung vào file chung
    with open(output_file, "w", encoding="utf-8") as f:
        f.writelines(full_text)
    print(f"Đã xử lý tất cả file và lưu vào {output_file}")

    # Xóa thư mục tạm
    if os.path.exists("temp"):
        shutil.rmtree("temp")

# Gọi hàm với thư mục chứa các file Word
directory = r"Data_track2_Word"
process_all_word_files(directory)